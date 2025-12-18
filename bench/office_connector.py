#!/usr/bin/env python3
"""
Office Suite Connector for Memex.

Pulls actual file content from Microsoft 365 or Google Workspace,
extracts knowledge using LLM, and ingests into memex.

This is NOT metadata-only - we get the actual text content to build
meaningful knowledge graphs.

Supported sources:
- Microsoft 365: OneDrive, SharePoint, Outlook (email body), OneNote
- Google Workspace: Drive (Docs, Sheets, Slides), Gmail (email body)

Usage:
    # Microsoft 365
    python office_connector.py --provider microsoft --auth-file ~/.memex/ms_credentials.json

    # Google Workspace
    python office_connector.py --provider google --auth-file ~/.memex/google_credentials.json

    # Sync specific folder
    python office_connector.py --provider google --folder "Shared Documents"

Setup:
    1. Create app registration in Azure AD or Google Cloud Console
    2. Download credentials JSON
    3. Run with --setup to authorize
"""

import argparse
import asyncio
import hashlib
import json
import logging
import mimetypes
import os
import re
import sys
import tempfile
import time
from abc import ABC, abstractmethod
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, AsyncIterator, Optional

import httpx
from dotenv import load_dotenv
from openai import AsyncOpenAI

# Optional imports for document parsing
try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

load_dotenv()

MEMEX_URL = os.getenv("MEMEX_URL", "http://localhost:8080")
MODEL = os.getenv("EXTRACTION_MODEL", "gpt-4o-mini")

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
)
logger = logging.getLogger(__name__)

# Extraction prompt (same as ingest_ai.py for consistency)
EXTRACTION_PROMPT = """Analyze this document and extract structured knowledge.

Document Title: {title}
Document Type: {doc_type}
Source: {source}

Content:
{content}

Extract:
1. Named entities (people, places, organizations, concepts, projects, tools)
2. Relationships between entities
3. Key facts, decisions, or action items

Return JSON only, no markdown:
{{
  "entities": [
    {{"name": "Full Name", "type": "Person|Organization|Project|Tool|Concept|Location|Event", "description": "brief description"}}
  ],
  "relationships": [
    {{"source": "Entity Name", "target": "Entity Name", "type": "RELATIONSHIP_TYPE", "properties": {{}}}}
  ],
  "summary": "1-2 sentence summary of document content"
}}

Relationship types: WORKS_ON, MANAGES, CREATED, OWNS, DEPENDS_ON, RELATED_TO, ASSIGNED_TO, DECIDED, MEMBER_OF, etc.

Focus on work-relevant knowledge. Extract what's explicitly stated."""


@dataclass
class Document:
    """A document from an office suite."""
    id: str
    name: str
    mime_type: str
    content: str  # Actual text content
    source: str  # "microsoft" or "google"
    path: str  # Full path in drive
    modified_time: Optional[str] = None
    author: Optional[str] = None
    url: Optional[str] = None
    meta: Optional[dict] = None


class OfficeProvider(ABC):
    """Base class for office suite providers."""

    @abstractmethod
    async def authenticate(self, credentials_path: str) -> bool:
        """Authenticate with the provider."""
        pass

    @abstractmethod
    async def list_files(
        self,
        folder: Optional[str] = None,
        file_types: Optional[list[str]] = None,
        modified_after: Optional[datetime] = None,
    ) -> AsyncIterator[dict]:
        """List files from the provider."""
        pass

    @abstractmethod
    async def get_file_content(self, file_id: str, mime_type: str) -> str:
        """Get the actual text content of a file."""
        pass

    @abstractmethod
    async def get_emails(
        self,
        folder: str = "inbox",
        limit: int = 100,
        modified_after: Optional[datetime] = None,
    ) -> AsyncIterator[dict]:
        """Get emails with body content."""
        pass


class MicrosoftProvider(OfficeProvider):
    """Microsoft 365 provider using Graph API."""

    GRAPH_URL = "https://graph.microsoft.com/v1.0"

    def __init__(self, admin_mode: bool = False):
        """
        Initialize Microsoft provider.

        Args:
            admin_mode: If True, use application permissions to access all users.
                       Requires admin consent in Azure AD.
        """
        self.admin_mode = admin_mode
        self.access_token: Optional[str] = None
        self.http_client: Optional[httpx.AsyncClient] = None
        self.users: list[dict] = []  # Populated in admin mode

    async def authenticate(self, credentials_path: str) -> bool:
        """
        Authenticate with Microsoft Graph API.

        Credentials file should contain:
        {
            "client_id": "...",
            "client_secret": "...",  # Required for admin/app-only flow
            "tenant_id": "...",      # Required for admin flow
            "access_token": "..."    # Or pre-obtained token for user flow
        }

        For admin mode, the app must have Application permissions:
        - Files.Read.All, Mail.Read, User.Read.All
        And admin must grant consent in Azure AD.
        """
        try:
            with open(credentials_path) as f:
                creds = json.load(f)

            # If we have a pre-obtained access token, use it
            if "access_token" in creds and not self.admin_mode:
                self.access_token = creds["access_token"]
            else:
                # Client credentials flow (app-only / admin)
                if "tenant_id" not in creds or "client_secret" not in creds:
                    logger.error("Admin mode requires tenant_id and client_secret")
                    return False

                async with httpx.AsyncClient() as client:
                    token_url = f"https://login.microsoftonline.com/{creds['tenant_id']}/oauth2/v2.0/token"
                    resp = await client.post(
                        token_url,
                        data={
                            "client_id": creds["client_id"],
                            "client_secret": creds["client_secret"],
                            "scope": "https://graph.microsoft.com/.default",
                            "grant_type": "client_credentials",
                        },
                    )
                    resp.raise_for_status()
                    self.access_token = resp.json()["access_token"]

            # Create authenticated client
            self.http_client = httpx.AsyncClient(
                base_url=self.GRAPH_URL,
                headers={"Authorization": f"Bearer {self.access_token}"},
                timeout=60,
            )

            # Verify authentication
            if self.admin_mode:
                # In admin mode, list users to verify permissions
                resp = await self.http_client.get("/users?$top=5")
                if resp.status_code == 200:
                    logger.info("Microsoft Graph API authenticated (ADMIN MODE)")
                    return True
                else:
                    logger.error(f"Admin auth failed - ensure Application permissions are granted: {resp.status_code}")
                    return False
            else:
                resp = await self.http_client.get("/me")
                if resp.status_code in (200, 403):
                    logger.info("Microsoft Graph API authenticated successfully")
                    return True
                else:
                    logger.error(f"Auth verification failed: {resp.status_code}")
                    return False

        except Exception as e:
            logger.error(f"Microsoft authentication failed: {e}")
            return False

    async def list_users(self) -> list[dict]:
        """List all users in the organization (admin mode only)."""
        if not self.admin_mode or not self.http_client:
            return []

        users = []
        url = "/users?$select=id,displayName,mail,userPrincipalName"

        while url:
            try:
                resp = await self.http_client.get(url)
                if resp.status_code != 200:
                    break
                data = resp.json()
                users.extend(data.get("value", []))
                url = data.get("@odata.nextLink")
            except Exception as e:
                logger.error(f"Error listing users: {e}")
                break

        self.users = users
        logger.info(f"Found {len(users)} users in organization")
        return users

    async def list_files(
        self,
        folder: Optional[str] = None,
        file_types: Optional[list[str]] = None,
        modified_after: Optional[datetime] = None,
        user_id: Optional[str] = None,
    ) -> AsyncIterator[dict]:
        """List files from OneDrive/SharePoint."""
        if not self.http_client:
            return

        # In admin mode, iterate over all users if no specific user
        if self.admin_mode and not user_id:
            if not self.users:
                await self.list_users()

            for user in self.users:
                uid = user.get("id")
                user_name = user.get("displayName", "Unknown")
                logger.info(f"Scanning files for user: {user_name}")

                async for file_info in self.list_files(
                    folder=folder,
                    file_types=file_types,
                    modified_after=modified_after,
                    user_id=uid,
                ):
                    file_info["owner"] = user_name
                    file_info["owner_email"] = user.get("mail") or user.get("userPrincipalName")
                    yield file_info
            return

        # Build base path - use user ID in admin mode, /me for user mode
        if user_id:
            base = f"/users/{user_id}/drive"
        else:
            base = "/me/drive"

        # Start from root or specific folder
        if folder:
            url = f"{base}/root:/{folder}:/children"
        else:
            url = f"{base}/root/children"

        params = {"$top": 100}

        if modified_after:
            params["$filter"] = f"lastModifiedDateTime ge {modified_after.isoformat()}Z"

        while url:
            try:
                resp = await self.http_client.get(url, params=params)
                if resp.status_code != 200:
                    logger.error(f"Failed to list files: {resp.status_code}")
                    break

                data = resp.json()

                for item in data.get("value", []):
                    if "file" in item:  # It's a file, not folder
                        mime_type = item.get("file", {}).get("mimeType", "")

                        # Filter by type if specified
                        if file_types:
                            if not any(ft in mime_type for ft in file_types):
                                continue

                        yield {
                            "id": item["id"],
                            "name": item["name"],
                            "mime_type": mime_type,
                            "path": item.get("parentReference", {}).get("path", "") + "/" + item["name"],
                            "modified_time": item.get("lastModifiedDateTime"),
                            "author": item.get("lastModifiedBy", {}).get("user", {}).get("displayName"),
                            "url": item.get("webUrl"),
                            "size": item.get("size", 0),
                        }
                    elif "folder" in item:
                        # Recursively list folder contents
                        async for sub_item in self.list_files(
                            folder=item["name"] if not folder else f"{folder}/{item['name']}",
                            file_types=file_types,
                            modified_after=modified_after,
                            user_id=user_id,
                        ):
                            yield sub_item

                # Handle pagination
                url = data.get("@odata.nextLink")
                params = {}  # Next link includes params

            except Exception as e:
                logger.error(f"Error listing files: {e}")
                break

    async def get_file_content(self, file_id: str, mime_type: str) -> str:
        """Get actual text content from a file."""
        if not self.http_client:
            return ""

        try:
            # For Office documents, use the content endpoint with format conversion
            if "word" in mime_type or "document" in mime_type:
                # Get Word document content
                resp = await self.http_client.get(
                    f"/me/drive/items/{file_id}/content",
                )
                if resp.status_code == 200:
                    return self._extract_docx(resp.content)

            elif "spreadsheet" in mime_type or "excel" in mime_type:
                # Get Excel content
                resp = await self.http_client.get(
                    f"/me/drive/items/{file_id}/content",
                )
                if resp.status_code == 200:
                    return self._extract_xlsx(resp.content)

            elif "presentation" in mime_type or "powerpoint" in mime_type:
                # PowerPoint - get as PDF and extract
                resp = await self.http_client.get(
                    f"/me/drive/items/{file_id}/content?format=pdf",
                )
                if resp.status_code == 200:
                    return self._extract_pdf(resp.content)

            elif "pdf" in mime_type:
                resp = await self.http_client.get(
                    f"/me/drive/items/{file_id}/content",
                )
                if resp.status_code == 200:
                    return self._extract_pdf(resp.content)

            elif "text" in mime_type or mime_type == "":
                # Plain text
                resp = await self.http_client.get(
                    f"/me/drive/items/{file_id}/content",
                )
                if resp.status_code == 200:
                    return resp.text

            else:
                logger.warning(f"Unsupported mime type: {mime_type}")
                return ""

        except Exception as e:
            logger.error(f"Error getting file content: {e}")
            return ""

        return ""

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX."""
        if not HAS_DOCX:
            logger.warning("python-docx not installed. Install with: pip install python-docx")
            return ""
        try:
            doc = DocxDocument(BytesIO(content))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return ""

    def _extract_xlsx(self, content: bytes) -> str:
        """Extract text from Excel."""
        if not HAS_XLSX:
            logger.warning("openpyxl not installed. Install with: pip install openpyxl")
            return ""
        try:
            wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
            text_parts = []
            for sheet in wb.worksheets:
                text_parts.append(f"=== Sheet: {sheet.title} ===")
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join(str(cell) if cell else "" for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting XLSX: {e}")
            return ""

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF."""
        if not HAS_PYPDF:
            logger.warning("pypdf not installed. Install with: pip install pypdf")
            return ""
        try:
            reader = PdfReader(BytesIO(content))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            return ""

    async def get_emails(
        self,
        folder: str = "inbox",
        limit: int = 100,
        modified_after: Optional[datetime] = None,
    ) -> AsyncIterator[dict]:
        """Get emails with body content from Outlook."""
        if not self.http_client:
            return

        params = {
            "$top": min(limit, 50),
            "$select": "id,subject,from,toRecipients,body,receivedDateTime,webLink",
        }

        if modified_after:
            params["$filter"] = f"receivedDateTime ge {modified_after.isoformat()}Z"

        url = f"/me/mailFolders/{folder}/messages"
        count = 0

        while url and count < limit:
            try:
                resp = await self.http_client.get(url, params=params)
                if resp.status_code != 200:
                    logger.error(f"Failed to get emails: {resp.status_code}")
                    break

                data = resp.json()

                for msg in data.get("value", []):
                    if count >= limit:
                        break

                    # Get actual email body text
                    body = msg.get("body", {})
                    content = body.get("content", "")

                    # Strip HTML if HTML format
                    if body.get("contentType") == "html":
                        content = self._strip_html(content)

                    yield {
                        "id": msg["id"],
                        "name": msg.get("subject", "No Subject"),
                        "mime_type": "message/rfc822",
                        "content": content,
                        "path": f"/{folder}/{msg.get('subject', '')}",
                        "modified_time": msg.get("receivedDateTime"),
                        "author": msg.get("from", {}).get("emailAddress", {}).get("name"),
                        "url": msg.get("webLink"),
                        "meta": {
                            "from": msg.get("from", {}).get("emailAddress", {}),
                            "to": [r.get("emailAddress", {}) for r in msg.get("toRecipients", [])],
                        },
                    }
                    count += 1

                url = data.get("@odata.nextLink")
                params = {}

            except Exception as e:
                logger.error(f"Error getting emails: {e}")
                break

    def _strip_html(self, html: str) -> str:
        """Strip HTML tags and decode entities."""
        import html as html_module
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', ' ', html)
        # Decode entities
        text = html_module.unescape(text)
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        return text


class GoogleProvider(OfficeProvider):
    """Google Workspace provider using Drive and Gmail APIs."""

    DRIVE_URL = "https://www.googleapis.com/drive/v3"
    GMAIL_URL = "https://gmail.googleapis.com/gmail/v1"
    DOCS_URL = "https://docs.googleapis.com/v1"
    ADMIN_URL = "https://admin.googleapis.com/admin/directory/v1"

    def __init__(self, admin_mode: bool = False):
        """
        Initialize Google provider.

        Args:
            admin_mode: If True, use domain-wide delegation to access all users.
                       Requires a service account with domain-wide delegation enabled.
        """
        self.admin_mode = admin_mode
        self.access_token: Optional[str] = None
        self.http_client: Optional[httpx.AsyncClient] = None
        self.service_account_creds: Optional[dict] = None
        self.users: list[dict] = []
        self.current_user_email: Optional[str] = None  # For impersonation

    async def authenticate(self, credentials_path: str) -> bool:
        """
        Authenticate with Google APIs.

        For user mode - credentials file should contain:
        {
            "access_token": "...",
            "refresh_token": "...",
            "client_id": "...",
            "client_secret": "..."
        }

        For admin mode - use service account JSON with domain-wide delegation:
        {
            "type": "service_account",
            "project_id": "...",
            "private_key": "...",
            "client_email": "...",
            "admin_email": "admin@domain.com"  # Admin to impersonate for directory API
        }
        """
        try:
            with open(credentials_path) as f:
                creds = json.load(f)

            if self.admin_mode:
                # Service account with domain-wide delegation
                if creds.get("type") != "service_account":
                    logger.error("Admin mode requires a service account JSON file")
                    return False

                self.service_account_creds = creds

                # Get token impersonating admin for directory access
                admin_email = creds.get("admin_email")
                if not admin_email:
                    logger.error("admin_email required in credentials for admin mode")
                    return False

                token = await self._get_service_account_token(
                    creds,
                    scopes=[
                        "https://www.googleapis.com/auth/admin.directory.user.readonly",
                        "https://www.googleapis.com/auth/drive.readonly",
                        "https://www.googleapis.com/auth/gmail.readonly",
                    ],
                    impersonate_user=admin_email,
                )
                if not token:
                    return False

                self.access_token = token
                self.current_user_email = admin_email

                # Create authenticated client
                self.http_client = httpx.AsyncClient(
                    headers={"Authorization": f"Bearer {self.access_token}"},
                    timeout=60,
                )

                logger.info("Google authenticated (ADMIN MODE - Domain-wide delegation)")
                return True

            else:
                # Regular OAuth flow
                if "access_token" in creds:
                    self.access_token = creds["access_token"]
                elif "refresh_token" in creds:
                    async with httpx.AsyncClient() as client:
                        resp = await client.post(
                            "https://oauth2.googleapis.com/token",
                            data={
                                "client_id": creds["client_id"],
                                "client_secret": creds["client_secret"],
                                "refresh_token": creds["refresh_token"],
                                "grant_type": "refresh_token",
                            },
                        )
                        resp.raise_for_status()
                        self.access_token = resp.json()["access_token"]
                else:
                    logger.error("No access_token or refresh_token in credentials")
                    return False

                # Create authenticated client
                self.http_client = httpx.AsyncClient(
                    headers={"Authorization": f"Bearer {self.access_token}"},
                    timeout=60,
                )

                # Verify
                resp = await self.http_client.get(f"{self.DRIVE_URL}/about?fields=user")
                if resp.status_code == 200:
                    user = resp.json().get("user", {})
                    logger.info(f"Google authenticated as: {user.get('displayName', 'Unknown')}")
                    return True
                else:
                    logger.error(f"Auth verification failed: {resp.status_code}")
                    return False

        except Exception as e:
            logger.error(f"Google authentication failed: {e}")
            return False

    async def _get_service_account_token(
        self,
        creds: dict,
        scopes: list[str],
        impersonate_user: Optional[str] = None,
    ) -> Optional[str]:
        """Get access token using service account credentials."""
        try:
            import jwt
            import time

            now = int(time.time())

            claim_set = {
                "iss": creds["client_email"],
                "scope": " ".join(scopes),
                "aud": "https://oauth2.googleapis.com/token",
                "iat": now,
                "exp": now + 3600,
            }

            if impersonate_user:
                claim_set["sub"] = impersonate_user

            # Sign JWT
            signed_jwt = jwt.encode(
                claim_set,
                creds["private_key"],
                algorithm="RS256",
            )

            # Exchange for access token
            async with httpx.AsyncClient() as client:
                resp = await client.post(
                    "https://oauth2.googleapis.com/token",
                    data={
                        "grant_type": "urn:ietf:params:oauth:grant-type:jwt-bearer",
                        "assertion": signed_jwt,
                    },
                )
                resp.raise_for_status()
                return resp.json()["access_token"]

        except ImportError:
            logger.error("PyJWT required for service account auth: pip install PyJWT")
            return None
        except Exception as e:
            logger.error(f"Service account token error: {e}")
            return None

    async def list_users(self) -> list[dict]:
        """List all users in the Google Workspace domain (admin mode only)."""
        if not self.admin_mode or not self.http_client:
            return []

        users = []
        url = f"{self.ADMIN_URL}/users?customer=my_customer&maxResults=500"

        while url:
            try:
                resp = await self.http_client.get(url)
                if resp.status_code != 200:
                    logger.error(f"Failed to list users: {resp.status_code} - {resp.text}")
                    break
                data = resp.json()
                users.extend(data.get("users", []))

                next_page = data.get("nextPageToken")
                if next_page:
                    url = f"{self.ADMIN_URL}/users?customer=my_customer&maxResults=500&pageToken={next_page}"
                else:
                    url = None
            except Exception as e:
                logger.error(f"Error listing users: {e}")
                break

        self.users = users
        logger.info(f"Found {len(users)} users in Google Workspace domain")
        return users

    async def _switch_user(self, user_email: str) -> bool:
        """Switch to impersonate a different user (admin mode only)."""
        if not self.admin_mode or not self.service_account_creds:
            return False

        token = await self._get_service_account_token(
            self.service_account_creds,
            scopes=[
                "https://www.googleapis.com/auth/drive.readonly",
                "https://www.googleapis.com/auth/gmail.readonly",
            ],
            impersonate_user=user_email,
        )

        if token:
            self.access_token = token
            self.current_user_email = user_email
            self.http_client = httpx.AsyncClient(
                headers={"Authorization": f"Bearer {token}"},
                timeout=60,
            )
            return True
        return False

    async def list_files(
        self,
        folder: Optional[str] = None,
        file_types: Optional[list[str]] = None,
        modified_after: Optional[datetime] = None,
        user_email: Optional[str] = None,
    ) -> AsyncIterator[dict]:
        """List files from Google Drive."""
        if not self.http_client:
            return

        # In admin mode, iterate over all users if no specific user
        if self.admin_mode and not user_email:
            if not self.users:
                await self.list_users()

            for user in self.users:
                email = user.get("primaryEmail")
                user_name = user.get("name", {}).get("fullName", email)
                logger.info(f"Scanning files for user: {user_name}")

                # Switch to impersonate this user
                if not await self._switch_user(email):
                    logger.warning(f"Could not impersonate {email}, skipping")
                    continue

                async for file_info in self.list_files(
                    folder=folder,
                    file_types=file_types,
                    modified_after=modified_after,
                    user_email=email,
                ):
                    file_info["owner"] = user_name
                    file_info["owner_email"] = email
                    yield file_info
            return

        # Build query
        q_parts = ["trashed=false"]

        if folder:
            # Find folder ID first
            folder_resp = await self.http_client.get(
                f"{self.DRIVE_URL}/files",
                params={
                    "q": f"name='{folder}' and mimeType='application/vnd.google-apps.folder'",
                    "fields": "files(id)",
                },
            )
            if folder_resp.status_code == 200:
                folders = folder_resp.json().get("files", [])
                if folders:
                    q_parts.append(f"'{folders[0]['id']}' in parents")

        if modified_after:
            q_parts.append(f"modifiedTime > '{modified_after.isoformat()}'")

        # Exclude folders from results
        q_parts.append("mimeType != 'application/vnd.google-apps.folder'")

        params = {
            "q": " and ".join(q_parts),
            "fields": "nextPageToken,files(id,name,mimeType,modifiedTime,owners,webViewLink,parents)",
            "pageSize": 100,
        }

        page_token = None

        while True:
            if page_token:
                params["pageToken"] = page_token

            try:
                resp = await self.http_client.get(f"{self.DRIVE_URL}/files", params=params)
                if resp.status_code != 200:
                    logger.error(f"Failed to list files: {resp.status_code} - {resp.text}")
                    break

                data = resp.json()

                for item in data.get("files", []):
                    mime_type = item.get("mimeType", "")

                    # Filter by type if specified
                    if file_types:
                        if not any(ft in mime_type for ft in file_types):
                            continue

                    yield {
                        "id": item["id"],
                        "name": item["name"],
                        "mime_type": mime_type,
                        "path": "/" + item["name"],  # Simplified path
                        "modified_time": item.get("modifiedTime"),
                        "author": item.get("owners", [{}])[0].get("displayName"),
                        "url": item.get("webViewLink"),
                    }

                page_token = data.get("nextPageToken")
                if not page_token:
                    break

            except Exception as e:
                logger.error(f"Error listing files: {e}")
                break

    async def get_file_content(self, file_id: str, mime_type: str) -> str:
        """Get actual text content from a Google file."""
        if not self.http_client:
            return ""

        try:
            # Google Docs - export as plain text
            if mime_type == "application/vnd.google-apps.document":
                resp = await self.http_client.get(
                    f"{self.DRIVE_URL}/files/{file_id}/export",
                    params={"mimeType": "text/plain"},
                )
                if resp.status_code == 200:
                    return resp.text

            # Google Sheets - export as CSV
            elif mime_type == "application/vnd.google-apps.spreadsheet":
                resp = await self.http_client.get(
                    f"{self.DRIVE_URL}/files/{file_id}/export",
                    params={"mimeType": "text/csv"},
                )
                if resp.status_code == 200:
                    return resp.text

            # Google Slides - export as plain text
            elif mime_type == "application/vnd.google-apps.presentation":
                resp = await self.http_client.get(
                    f"{self.DRIVE_URL}/files/{file_id}/export",
                    params={"mimeType": "text/plain"},
                )
                if resp.status_code == 200:
                    return resp.text

            # Regular files - download and parse
            elif "pdf" in mime_type:
                resp = await self.http_client.get(
                    f"{self.DRIVE_URL}/files/{file_id}",
                    params={"alt": "media"},
                )
                if resp.status_code == 200:
                    return self._extract_pdf(resp.content)

            elif "word" in mime_type or "document" in mime_type:
                resp = await self.http_client.get(
                    f"{self.DRIVE_URL}/files/{file_id}",
                    params={"alt": "media"},
                )
                if resp.status_code == 200:
                    return self._extract_docx(resp.content)

            elif "text" in mime_type:
                resp = await self.http_client.get(
                    f"{self.DRIVE_URL}/files/{file_id}",
                    params={"alt": "media"},
                )
                if resp.status_code == 200:
                    return resp.text

            else:
                logger.warning(f"Unsupported mime type for content extraction: {mime_type}")
                return ""

        except Exception as e:
            logger.error(f"Error getting file content: {e}")
            return ""

        return ""

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF."""
        if not HAS_PYPDF:
            logger.warning("pypdf not installed")
            return ""
        try:
            reader = PdfReader(BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            return ""

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX."""
        if not HAS_DOCX:
            logger.warning("python-docx not installed")
            return ""
        try:
            doc = DocxDocument(BytesIO(content))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return ""

    async def get_emails(
        self,
        folder: str = "INBOX",
        limit: int = 100,
        modified_after: Optional[datetime] = None,
    ) -> AsyncIterator[dict]:
        """Get emails with body content from Gmail."""
        if not self.http_client:
            return

        # Build query
        q_parts = []
        if modified_after:
            after_ts = int(modified_after.timestamp())
            q_parts.append(f"after:{after_ts}")

        params = {
            "maxResults": min(limit, 100),
            "labelIds": folder,
        }
        if q_parts:
            params["q"] = " ".join(q_parts)

        count = 0

        try:
            resp = await self.http_client.get(
                f"{self.GMAIL_URL}/users/me/messages",
                params=params,
            )
            if resp.status_code != 200:
                logger.error(f"Failed to list emails: {resp.status_code}")
                return

            messages = resp.json().get("messages", [])

            for msg_ref in messages:
                if count >= limit:
                    break

                # Get full message
                msg_resp = await self.http_client.get(
                    f"{self.GMAIL_URL}/users/me/messages/{msg_ref['id']}",
                    params={"format": "full"},
                )
                if msg_resp.status_code != 200:
                    continue

                msg = msg_resp.json()

                # Extract headers
                headers = {h["name"]: h["value"] for h in msg.get("payload", {}).get("headers", [])}

                # Extract body
                content = self._extract_gmail_body(msg.get("payload", {}))

                yield {
                    "id": msg["id"],
                    "name": headers.get("Subject", "No Subject"),
                    "mime_type": "message/rfc822",
                    "content": content,
                    "path": f"/{folder}/{headers.get('Subject', '')}",
                    "modified_time": datetime.fromtimestamp(int(msg.get("internalDate", 0)) / 1000).isoformat(),
                    "author": headers.get("From"),
                    "meta": {
                        "from": headers.get("From"),
                        "to": headers.get("To"),
                        "date": headers.get("Date"),
                    },
                }
                count += 1

        except Exception as e:
            logger.error(f"Error getting emails: {e}")

    def _extract_gmail_body(self, payload: dict) -> str:
        """Extract body text from Gmail message payload."""
        import base64

        # Check for direct body
        body = payload.get("body", {})
        if body.get("data"):
            try:
                return base64.urlsafe_b64decode(body["data"]).decode("utf-8")
            except:
                pass

        # Check parts
        parts = payload.get("parts", [])
        for part in parts:
            mime_type = part.get("mimeType", "")
            if mime_type == "text/plain":
                data = part.get("body", {}).get("data")
                if data:
                    try:
                        return base64.urlsafe_b64decode(data).decode("utf-8")
                    except:
                        pass
            elif mime_type == "text/html":
                data = part.get("body", {}).get("data")
                if data:
                    try:
                        html = base64.urlsafe_b64decode(data).decode("utf-8")
                        return self._strip_html(html)
                    except:
                        pass
            elif "multipart" in mime_type:
                # Recursive
                result = self._extract_gmail_body(part)
                if result:
                    return result

        return ""

    def _strip_html(self, html: str) -> str:
        """Strip HTML tags."""
        import html as html_module
        text = re.sub(r'<[^>]+>', ' ', html)
        text = html_module.unescape(text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text


class OfficeConnector:
    """Main connector that ingests office content into Memex."""

    def __init__(
        self,
        provider: OfficeProvider,
        memex_url: str = MEMEX_URL,
        llm_model: str = MODEL,
    ):
        self.provider = provider
        self.memex_url = memex_url
        self.llm_model = llm_model
        self.llm_client: Optional[AsyncOpenAI] = None
        self.http_client: Optional[httpx.AsyncClient] = None

        # Progress tracking
        self.processed_ids: set[str] = set()
        self.stats = {
            "files_processed": 0,
            "entities_created": 0,
            "relationships_created": 0,
            "errors": 0,
        }

    async def setup(self):
        """Initialize clients."""
        self.llm_client = AsyncOpenAI()
        self.http_client = httpx.AsyncClient(base_url=self.memex_url, timeout=60)

    async def cleanup(self):
        """Clean up resources."""
        if self.http_client:
            await self.http_client.aclose()

    def make_source_id(self, doc: Document) -> str:
        """Create content-addressed source ID."""
        content_hash = hashlib.sha256((doc.id + doc.content[:1000]).encode()).hexdigest()[:16]
        return f"source:office:{content_hash}"

    def make_entity_id(self, name: str, entity_type: str) -> str:
        """Create deterministic entity ID."""
        normalized = re.sub(r'[^a-z0-9]', '-', name.lower().strip())
        normalized = re.sub(r'-+', '-', normalized).strip('-')
        return f"{entity_type.lower()}:{normalized}"

    async def extract_knowledge(self, doc: Document) -> Optional[dict]:
        """Use LLM to extract entities and relationships."""
        if not self.llm_client:
            return None

        # Truncate content if too long
        content = doc.content[:15000]

        try:
            response = await self.llm_client.chat.completions.create(
                model=self.llm_model,
                messages=[{
                    "role": "user",
                    "content": EXTRACTION_PROMPT.format(
                        title=doc.name,
                        doc_type=doc.mime_type,
                        source=doc.source,
                        content=content,
                    ),
                }],
                max_tokens=4000,
            )

            text = response.choices[0].message.content
            if not text:
                return None

            text = text.strip()

            # Handle markdown code blocks
            if "```" in text:
                match = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
                if match:
                    text = match.group(1)

            return json.loads(text)

        except json.JSONDecodeError as e:
            logger.warning(f"JSON parse error for '{doc.name}': {e}")
        except Exception as e:
            logger.warning(f"Extraction error for '{doc.name}': {e}")

        return None

    async def ingest_document(self, doc: Document) -> bool:
        """Ingest a single document into Memex."""
        if not self.http_client:
            return False

        source_id = self.make_source_id(doc)

        # Skip if already processed
        if source_id in self.processed_ids:
            return True

        # 1. Create source node with actual content
        source_node = {
            "id": source_id,
            "type": "Source",
            "content": doc.content[:50000],  # Limit stored content
            "meta": {
                "title": doc.name,
                "path": doc.path,
                "source_type": "office",
                "provider": doc.source,
                "mime_type": doc.mime_type,
                "modified_time": doc.modified_time,
                "author": doc.author,
                "url": doc.url,
            },
        }

        try:
            resp = await self.http_client.post("/api/nodes", json=source_node)
            if resp.status_code not in (200, 201):
                logger.debug(f"Failed to create source node: {resp.status_code}")
        except Exception as e:
            logger.debug(f"Error creating source node: {e}")

        # 2. Extract knowledge using LLM
        extraction = await self.extract_knowledge(doc)

        if not extraction:
            self.processed_ids.add(source_id)
            self.stats["files_processed"] += 1
            self.stats["errors"] += 1
            return False

        entity_ids = []

        # 3. Create entity nodes
        for entity in extraction.get("entities", []):
            try:
                entity_id = self.make_entity_id(entity["name"], entity["type"])
                entity_ids.append(entity_id)

                entity_node = {
                    "id": entity_id,
                    "type": entity["type"],
                    "content": entity.get("description", entity["name"]),
                    "meta": {
                        "name": entity["name"],
                        "extracted_from": source_id,
                    },
                }

                resp = await self.http_client.post("/api/nodes", json=entity_node)
                if resp.status_code in (200, 201):
                    self.stats["entities_created"] += 1

                # Link to source
                await self.http_client.post("/api/links", json={
                    "source": entity_id,
                    "target": source_id,
                    "type": "EXTRACTED_FROM",
                    "meta": {},
                })

            except Exception as e:
                logger.debug(f"Error creating entity: {e}")

        # 4. Create relationship edges
        for rel in extraction.get("relationships", []):
            try:
                source_entity_id = self.make_entity_id(rel["source"], "Entity")
                target_entity_id = self.make_entity_id(rel["target"], "Entity")

                # Try to find actual IDs
                for eid in entity_ids:
                    if rel["source"].lower().replace(" ", "-") in eid:
                        source_entity_id = eid
                    if rel["target"].lower().replace(" ", "-") in eid:
                        target_entity_id = eid

                resp = await self.http_client.post("/api/links", json={
                    "source": source_entity_id,
                    "target": target_entity_id,
                    "type": rel["type"],
                    "meta": rel.get("properties", {}),
                })
                if resp.status_code in (200, 201):
                    self.stats["relationships_created"] += 1

            except Exception as e:
                logger.debug(f"Error creating relationship: {e}")

        # 5. Update attention DAG for co-occurring entities
        for i, src in enumerate(entity_ids):
            for tgt in entity_ids[i + 1:]:
                try:
                    await self.http_client.post("/api/edges/attention", json={
                        "source": src,
                        "target": tgt,
                        "query_id": f"office_ingest:{source_id}",
                        "weight": 0.7,
                    })
                except:
                    pass

        self.processed_ids.add(source_id)
        self.stats["files_processed"] += 1

        return True

    async def sync_files(
        self,
        folder: Optional[str] = None,
        file_types: Optional[list[str]] = None,
        limit: int = 100,
    ):
        """Sync files from office suite to Memex."""
        logger.info(f"Starting file sync (folder={folder}, limit={limit})")

        count = 0
        async for file_info in self.provider.list_files(folder=folder, file_types=file_types):
            if count >= limit:
                break

            logger.info(f"Processing: {file_info['name']}")

            # Get actual content
            content = await self.provider.get_file_content(
                file_info["id"],
                file_info["mime_type"],
            )

            if not content or len(content.strip()) < 50:
                logger.warning(f"Skipping {file_info['name']} - no content extracted")
                continue

            doc = Document(
                id=file_info["id"],
                name=file_info["name"],
                mime_type=file_info["mime_type"],
                content=content,
                source=self.provider.__class__.__name__.replace("Provider", "").lower(),
                path=file_info["path"],
                modified_time=file_info.get("modified_time"),
                author=file_info.get("author"),
                url=file_info.get("url"),
            )

            success = await self.ingest_document(doc)
            if success:
                logger.info(f"  Ingested: {len(content)} chars, "
                           f"entities: {self.stats['entities_created']}")

            count += 1

        logger.info(f"File sync complete: {self.stats}")

    async def sync_emails(
        self,
        folder: str = "inbox",
        limit: int = 100,
    ):
        """Sync emails to Memex."""
        logger.info(f"Starting email sync (folder={folder}, limit={limit})")

        count = 0
        async for email_info in self.provider.get_emails(folder=folder, limit=limit):
            if count >= limit:
                break

            logger.info(f"Processing email: {email_info['name'][:50]}...")

            content = email_info.get("content", "")
            if not content or len(content.strip()) < 20:
                continue

            doc = Document(
                id=email_info["id"],
                name=email_info["name"],
                mime_type="message/rfc822",
                content=content,
                source=self.provider.__class__.__name__.replace("Provider", "").lower(),
                path=email_info["path"],
                modified_time=email_info.get("modified_time"),
                author=email_info.get("author"),
                meta=email_info.get("meta"),
            )

            await self.ingest_document(doc)
            count += 1

        logger.info(f"Email sync complete: {self.stats}")


async def main():
    parser = argparse.ArgumentParser(description="Office Suite Connector for Memex")
    parser.add_argument(
        "--provider",
        choices=["microsoft", "google"],
        required=True,
        help="Office suite provider",
    )
    parser.add_argument(
        "--auth-file",
        required=True,
        help="Path to credentials JSON file",
    )
    parser.add_argument(
        "--admin",
        action="store_true",
        help="Admin mode: access all users in organization (requires admin consent)",
    )
    parser.add_argument(
        "--folder",
        help="Specific folder to sync",
    )
    parser.add_argument(
        "--sync-type",
        choices=["files", "emails", "all"],
        default="all",
        help="What to sync",
    )
    parser.add_argument(
        "--limit",
        type=int,
        default=100,
        help="Maximum items to process",
    )
    parser.add_argument(
        "--memex-url",
        default=MEMEX_URL,
        help="Memex API URL",
    )

    args = parser.parse_args()

    # Check for API key
    if not os.environ.get("OPENAI_API_KEY"):
        logger.error("OPENAI_API_KEY environment variable not set")
        return

    # Create provider
    if args.provider == "microsoft":
        provider = MicrosoftProvider(admin_mode=args.admin)
    else:
        provider = GoogleProvider(admin_mode=args.admin)

    # Authenticate
    logger.info(f"Authenticating with {args.provider}...")
    if not await provider.authenticate(args.auth_file):
        logger.error("Authentication failed")
        return

    # Create connector
    connector = OfficeConnector(provider, memex_url=args.memex_url)
    await connector.setup()

    try:
        if args.sync_type in ("files", "all"):
            await connector.sync_files(folder=args.folder, limit=args.limit)

        if args.sync_type in ("emails", "all"):
            await connector.sync_emails(limit=args.limit)

        logger.info("=" * 60)
        logger.info("OFFICE SYNC COMPLETE")
        logger.info("=" * 60)
        logger.info(f"Files processed: {connector.stats['files_processed']}")
        logger.info(f"Entities created: {connector.stats['entities_created']}")
        logger.info(f"Relationships created: {connector.stats['relationships_created']}")
        logger.info(f"Errors: {connector.stats['errors']}")

    finally:
        await connector.cleanup()


if __name__ == "__main__":
    asyncio.run(main())
